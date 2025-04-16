from langchain.memory import ConversationBufferMemory
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
import os
import json
import traceback
from dotenv import load_dotenv

load_dotenv()
class CustomerServiceAI:
    def __init__(self, llm, qa_data=None, qa_file=None, model="gpt-3.5-turbo" ):
        self.llm = llm
        self.qa_data = qa_data or []
        self.memory = ConversationBufferMemory()
        self.processing_status = {"status": "idle", "message": ""}
        self.vector_index_built = False
        self.images = []
        self.api_key = os.environ.get("OPENAI_API_KEY")
        self.use_llm_refinement = False
        self.model = model

        # 如果提供了 Q&A 文件，則載入
        if qa_file:
            if isinstance(qa_file, str) and qa_file.endswith('.json'):
                self.load_json_file(qa_file)
            elif isinstance(qa_file, str) and qa_file.endswith('.docx'):
                self.load_word_file(qa_file)

    def load_json_file(self, file_path):
        """載入 JSON 格式的 Q&A 資料"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                self.qa_data = json.load(f)

            # 建立向量索引
            if self.qa_data:
                self._build_vector_index()

            return self.qa_data
        except Exception as e:
            print(f"載入 JSON 文件時出錯: {str(e)}")
            return []

    def append_qa_to_json(self, new_qa_pairs, json_file_path="customer_service_qa.json"):
        """
        將新的問答對添加到現有的 JSON 文件中

        參數:
        new_qa_pairs (list): 包含新問答對的列表，每個問答對是一個字典，包含 "question" 和 "answer" 鍵
        json_file_path (str): JSON 文件的路徑

        返回:
        bool: 操作是否成功
        """
        try:
            # 檢查文件是否存在
            if os.path.exists(json_file_path):
                # 讀取現有的 JSON 文件
                with open(json_file_path, 'r', encoding='utf-8') as file:
                    existing_qa_pairs = json.load(file)
            else:
                # 如果文件不存在，創建一個空列表
                existing_qa_pairs = []

            # 檢查新問答對是否已存在，避免重複
            added_count = 0
            for new_qa in new_qa_pairs:
                if not any(qa['question'] == new_qa['question'] for qa in existing_qa_pairs):
                    existing_qa_pairs.append(new_qa)
                    added_count += 1

            # 將更新後的問答對寫入 JSON 文件
            with open(json_file_path, 'w', encoding='utf-8') as file:
                json.dump(existing_qa_pairs, file, ensure_ascii=False, indent=4)

            return added_count
        except Exception as e:
            print(f"添加問答對到 JSON 文件時出錯: {str(e)}")
            return 0

    def load_word_file(self, file_path: str):
        """載入 Word 檔案並提取問答對"""
        try:
            self.processing_status = {"status": "processing", "message": "正在讀取 Word 檔案..."}

            # 檢查文件是否存在
            if not os.path.exists(file_path):
                self.processing_status = {"status": "error", "message": f"找不到文件: {file_path}"}
                return

            # 使用 python-docx 處理 Word 文件
            from docx import Document

            # 確保文件路徑是正確的
            print(f"嘗試打開文件: {file_path}")
            doc = Document(file_path)
            print("文件打開成功")

            # 提取問答對
            qa_data = []
            current_question = None
            current_answer = None

            self.processing_status = {"status": "processing", "message": "正在解析問答對..."}

            # 遍歷文檔的段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # 檢查是否是問題開始
                if text.startswith("Q：") or text.startswith("Q:"):
                    # 如果已有一個問答對，先保存
                    if current_question and current_answer:
                        qa_data.append({
                            "question": current_question,
                            "answer": current_answer
                        })

                    # 開始新的問答對
                    current_question = text[2:].strip()
                    current_answer = None

                # 檢查是否是回答開始
                elif text.startswith("A：") or text.startswith("A:"):
                    current_answer = text[2:].strip()

                # 如果不是問題或回答的開始，但有當前回答，則將文本添加到當前回答
                elif current_answer is not None:
                    current_answer += "\n" + text

            # 添加最後一個問答對
            if current_question and current_answer:
                qa_data.append({
                    "question": current_question,
                    "answer": current_answer
                })

            self.qa_data = qa_data
            self.processing_status = {"status": "completed", "message": f"已成功解析 {len(qa_data)} 個問答對"}

            # 如果有問答對，建立向量索引
            if self.qa_data:
                # 將解析出的問答對添加到 JSON 文件
                added_count = self.append_qa_to_json(self.qa_data)

                if added_count > 0:
                    print(f"成功添加 {added_count} 個新問答對到知識庫！")
                else:
                    print("沒有新的問答對被添加，可能是因為所有問答對已存在。")
                self._build_vector_index()

            return self.qa_data

        except Exception as e:
            error_msg = traceback.format_exc()
            print(f"解析 Word 檔案時出錯: {error_msg}")
            self.processing_status = {"status": "error", "message": f"解析 Word 檔案時出錯: {str(e)}"}
            return []

    def _build_vector_index(self):
        """建立向量索引"""
        try:
            self.processing_status = {"status": "processing", "message": "正在建立向量索引..."}

            # 初始化嵌入模型
            self.embeddings = OpenAIEmbeddings()

            # 準備文檔
            documents = []
            for qa in self.qa_data:
                # 將問題和答案分開嵌入，以提高匹配精度
                question_doc = Document(
                    page_content=qa['question'],
                    metadata={"question": qa["question"], "answer": qa["answer"], "type": "question"}
                )
                documents.append(question_doc)

                # 也可以選擇性地將答案加入索引
                answer_doc = Document(
                    page_content=f"問題: {qa['question']}\n答案: {qa['answer']}",
                    metadata={"question": qa["question"], "answer": qa["answer"], "type": "answer"}
                )
                documents.append(answer_doc)

            # 建立向量存儲
            self.vector_store = FAISS.from_documents(documents, self.embeddings)
            self.vector_index_built = True

            self.processing_status = {"status": "completed", "message": "成功建立向量索引"}
            print(f"成功建立向量索引，包含 {len(documents)} 個文檔")
        except Exception as e:
            error_msg = traceback.format_exc()
            print(f"建立向量索引時出錯: {error_msg}")
            self.processing_status = {"status": "error", "message": f"建立向量索引時出錯: {str(e)}"}

    def answer_question(self, question, debug_callback=None):
        """回答用戶問題"""
        def debug(message):
            if debug_callback:
                debug_callback(message)
            print(message)

        try:
            debug(f"處理問題: {question}")

            # 首先嘗試直接文本匹配
            exact_match = self._exact_match_search(question, debug)
            if exact_match:
                debug("使用直接文本匹配的結果")
                debug(self.use_llm_refinement)
                if self.use_llm_refinement:
                    refined_answer = self.refine_answer_with_llm(exact_match, question)
                    debug("答案已經過 LLM 修飾")
                    return refined_answer
                else:
                    debug("未使用 LLM 修飾")
                    return exact_match

            # 如果已建立向量索引，使用向量搜索
            if self.vector_index_built:
                debug(f"使用向量搜索回答問題: {question}")

                # 使用 similarity_search_with_score 獲取分數
                docs_and_scores = self.vector_store.similarity_search_with_score(question, k=5)

                # 打印搜索結果以便調試
                debug("找到的相關文檔:")
                relevant_docs = []
                for i, (doc, score) in enumerate(docs_and_scores):
                    debug_msg = f"文檔 {i+1}:\n內容: {doc.page_content}\n相似度分數: {score}"
                    debug(debug_msg)

                    # 收集所有文檔
                    relevant_docs.append((doc, score))

                # 按相似度排序
                relevant_docs.sort(key=lambda x: x[1])

                # 如果找到了文檔
                if relevant_docs:
                    # 使用最相關的文檔作為上下文
                    context = "\n\n".join([doc.page_content for doc, _ in relevant_docs[:3]])

                    debug(f"使用上下文:\n{context}")

                    # 使用 LLM 生成回答
                    from langchain_core.prompts import PromptTemplate
                    from langchain_core.output_parsers import StrOutputParser

                    template = """
                    你是一個專業的客服助手。請根據以下參考資料回答用戶的問題。
                    如果參考資料中有直接相關的答案，請使用該答案。
                    如果參考資料中沒有相關信息，請誠實地說你不知道，不要編造答案。

                    參考資料:
                    {context}

                    用戶問題: {question}

                    請提供專業、有禮貌且有幫助的回答:
                    """

                    prompt = PromptTemplate(
                        template=template,
                        input_variables=["context", "question"]
                    )

                    chain = prompt | self.llm | StrOutputParser()

                    response = chain.invoke({
                        "context": context,
                        "question": question
                    })

                    return response

            # 如果向量搜索失敗或未建立索引，使用備用方法
            debug("使用備用方法")
            return self._fallback_answer(question, debug)

        except Exception as e:
            error_msg = traceback.format_exc()
            debug(f"回答問題時出錯: {error_msg}")
            return f"很抱歉，處理您的問題時出現了錯誤。請稍後再試或聯繫人工客服。錯誤詳情: {str(e)}"

    def _exact_match_search(self, question, debug=print):
        """嘗試直接文本匹配"""
        question_lower = question.lower()
        debug(f"進行直接文本匹配: {question_lower}")

        for qa in self.qa_data:
            qa_question_lower = qa["question"].lower()

            # 完全匹配
            if question_lower == qa_question_lower:
                debug(f"找到完全匹配: {qa['question']}")
                return qa["answer"]

            # 部分匹配
            if question_lower in qa_question_lower or qa_question_lower in question_lower:
                debug(f"找到部分匹配: {qa['question']}")
                return qa["answer"]

            # 檢查問題的關鍵部分是否匹配
            keywords = ["顯示名字", "顯示名稱", "進場通知", "看不到名字", "看不到名稱"]
            for keyword in keywords:
                if keyword in question_lower and keyword in qa_question_lower:
                    debug(f"找到關鍵詞匹配: {qa['question']} (關鍵詞: {keyword})")
                    return qa["answer"]

        debug("沒有找到直接文本匹配")
        return None

    def _fallback_answer(self, question, debug=print):
        """當向量搜索失敗時的備用方法"""
        # 嘗試直接關鍵詞匹配
        debug("使用關鍵詞匹配方法")

        # 將問題轉換為關鍵詞集合
        question_words = set(question.lower().split())
        debug(f"問題關鍵詞: {question_words}")

        best_match = None
        highest_match_count = 0

        for qa in self.qa_data:
            # 將問題轉換為關鍵詞集合
            qa_words = set(qa["question"].lower().split())

            # 計算匹配的關鍵詞數量
            match_count = len(question_words.intersection(qa_words))

            if match_count > highest_match_count:
                highest_match_count = match_count
                best_match = qa
                debug(f"找到更好的匹配: '{qa['question']}', 匹配關鍵詞數量: {match_count}")

        # 如果匹配度足夠高
        if highest_match_count >= 2 and best_match:
            debug(f"使用關鍵詞匹配結果: '{best_match['question']}'")
            debug(f"匹配關鍵詞數量: {highest_match_count}")
            return best_match["answer"]
        else:
            debug(f"關鍵詞匹配不足: 最高匹配數 {highest_match_count}")

        # 如果沒有找到匹配的問答對，使用 LLM 生成通用回答
        debug("沒有找到匹配的問答對，使用 LLM 生成通用回答")
        from langchain_core.prompts import PromptTemplate
        from langchain_core.output_parsers import StrOutputParser

        template = """
        你是一個專業的客服助手。請回答用戶的問題。
        如果你不知道答案，請誠實地說你不知道，並建議用戶聯繫人工客服。

        用戶問題: {question}

        請提供專業、有禮貌且有幫助的回答:
        """

        prompt = PromptTemplate(
            template=template,
            input_variables=["question"]
        )

        chain = prompt | self.llm | StrOutputParser()

        debug("使用 LLM 生成回答")
        response = chain.invoke({
            "question": question
        })

        return response

    def _calculate_similarity(self, text1, text2):
        """計算兩個文本的相似度 (簡單版本)"""
        # 將文本轉換為集合
        set1 = set(text1.split())
        set2 = set(text2.split())

        # 計算 Jaccard 相似度
        intersection = len(set1.intersection(set2))
        union = len(set1.union(set2))

        if union == 0:
            return 0
        return intersection / union

    def get_qa_data(self):
        """獲取當前的 Q&A 資料"""
        return self.qa_data

    def get_images(self):
        """獲取從 Word 文件中提取的圖片"""
        return self.images

    def _exact_match_search2(self, question):
        """嘗試直接文本匹配"""
        question_lower = question.lower()

        for qa in self.qa_data:
            if question_lower in qa["question"].lower():
                print(f"找到直接文本匹配: {qa['question']}")
                return qa["answer"]

            # 檢查問題的關鍵部分是否匹配
            # 例如，"如何顯示名字" 可以匹配 "我看不到其他人進來的顯示名字"
            keywords = ["顯示名字", "顯示名稱", "進場通知", "看不到名字", "看不到名稱"]
            for keyword in keywords:
                if keyword in question_lower and keyword in qa["question"].lower():
                    print(f"找到關鍵詞匹配: {qa['question']} (關鍵詞: {keyword})")
                    return qa["answer"]

        return None

    def refine_answer_with_llm(self, original_answer, question):
        """
        使用 LLM 修飾答案，優化用詞和語氣

        參數:
        original_answer (str): 原始答案
        question (str): 用戶問題

        返回:
        str: 修飾後的答案
        """
        try:
            # 使用 OpenAI API
            import openai

            # 確保設置了 API 密鑰
            if not hasattr(self, 'openai_client'):
                self.openai_client = openai.OpenAI(api_key=self.api_key)

            # 構建提示
            prompt = f"""
            請優化以下客服回答，使其更專業、親切且易於理解。保持原始資訊完整，但改善用詞、語氣和結構。

            用戶問題: {question}

            原始回答:
            {original_answer}

            優化後的回答:
            """

            # 調用 LLM
            response = self.openai_client.chat.completions.create(
                model=self.model,  # 使用您設置的模型
                messages=[
                    {"role": "system", "content": "你是一位專業的客服優化專家，擅長將回答修飾得更加專業、親切且易於理解。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,  # 較低的溫度以保持一致性
                max_tokens=1000
            )

            # 獲取修飾後的答案
            refined_answer = response.choices[0].message.content.strip()

            return refined_answer
        except Exception as e:
            print(f"使用 LLM 修飾答案時出錯: {str(e)}")
            # 如果出錯，返回原始答案
            return original_answer
